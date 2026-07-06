from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def add_cover_page(doc,title):
    """
    Creates a professional academic cover page.
    """

    p=doc.add_paragraph()
    p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

    run=p.add_run("PROJECT REPORT\n\n")
    run.bold=True
    run.font.size=Pt(24)

    run=p.add_run(title.upper()+"\n\n")
    run.bold=True
    run.font.size=Pt(20)

    run=p.add_run("Generated using Athena AI\n")
    run.font.size=Pt(16)

    run=p.add_run("AI Engineering Assistant\n\n")
    run.font.size=Pt(14)

    run=p.add_run("Version 5.0")
    run.italic=True

    doc.add_page_break()